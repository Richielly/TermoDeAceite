from docx2pdf import convert
from docx import Document
# import pythoncom
# import win32com.client
import time
import os

class Arquivo():
    # xl=win32com.client.Dispatch("Excel.Application",pythoncom.CoInitialize())

    def alterar_doc(self, _cliente, _nrProcesso, _tipoProcesso, _anoProcesso, _contrato, _dia, _mes, _ano, _cidade, _rua, _numero, _sistema, _responsavel, _responsavelImplantacao, _responsavelAssinatura, _responsavelCargo, _responsavelNrDocumento):
        document = Document('PROC001_PROCESSO_DE_CONVERSAO_E_IMPLANTACAO.docx')

        cliente = _cliente
        contrato = _contrato
        dia = _dia
        mes = _mes
        ano = _ano
        cidade = _cidade
        rua = _rua
        numero = _numero
        sistema = _sistema
        responsavel = _responsavel
        tipoProcesso = _tipoProcesso
        nrProcesso = _nrProcesso
        anoProcesso = _anoProcesso
        responsavelImplantacao = _responsavelImplantacao
        responsavelAssinatura = _responsavelAssinatura
        responsavelCargo = _responsavelCargo
        responsavelNrDocumento = _responsavelNrDocumento

        parag_3 = document.paragraphs[3]
        parag_3.text = str(parag_3.text).replace('{$cliente}', cliente).replace('{$nrProcesso}',nrProcesso).replace('{$anoProcesso}', anoProcesso).replace('{$contrato}', contrato)

        parag_8 = document.paragraphs[8]
        parag_8.text = str(parag_8.text).replace('{$dia}',dia).replace('{$mes}',mes).replace('{$ano}',ano).replace('{$cidade}',cidade).replace('{$rua}',rua).replace('{$numero}',numero).replace('{$sistema}',sistema).replace('{$responsavel}',responsavel)

        parag_9 = document.paragraphs[9]
        parag_9.text = str(parag_9.text).replace('{$tipoProcesso}', tipoProcesso).replace('{$nrProcesso}', nrProcesso).replace('{$anoProcesso}', anoProcesso)

        parag_16 = document.paragraphs[16]
        parag_16.text = str(parag_16.text).replace('{$responsavelImplantacao}',responsavelImplantacao)

        parag_18 = document.paragraphs[18]
        parag_18.text = str(parag_18.text).replace('{$responsavelAssinatura}', responsavelAssinatura)

        parag_19 = document.paragraphs[19]
        parag_19.text = str(parag_19.text).replace('{$responsavelCargo}', responsavelCargo)

        parag_20 = document.paragraphs[20]
        parag_20.text = str(parag_20.text).replace('{$responsavelNrDocumento}', responsavelNrDocumento)

        unico = (str(time.time()).replace('.',''))

        document.save(f"""TermoScp_.docx""")

        return f'TermoScp_'

    def converter(self, nome):
        # xl = win32com.client.Dispatch("Excel.Application", pythoncom.CoInitialize())
        try:
            msg_dir = os.getcwd()
            dirs = os.listdir(msg_dir)
            inputFile = f"""{nome}.docx"""
            # outputFile = f"""{nome}.pdf"""
            # file = open(outputFile, "w")
            # file.close()
            msg = convert(inputFile)
            return dirs
        except:
            return  str(dirs)
            # return f"""Depois de clicar no botão abaixo, procure pelo arquivo {nome}'.pdf  na pasta de downloads."""